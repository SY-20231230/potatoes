import os
import pymysql
import json
from datetime import datetime, timedelta
from fastapi import FastAPI
from pydantic import BaseModel
import google.generativeai as genai
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from io import BytesIO
from collections import Counter

from docx import Document
from docx.shared import Pt, Inches

# 🌐 FastAPI 앱 초기화
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 프론트엔드 주소
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# 🔑 Gemini API 키 설정
GEMINI_API_KEY = "AIzaSyA0MXb6rhjiHsbgGJj9EbLi_oXg3iUU0vY"
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")

# 📩 질문 스키마
class Question(BaseModel):
    user_question: str

def get_geocache_data():
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password="yhkcar@!0293",
        db="jolupdb",
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor
    )
    try:
        with conn.cursor() as cursor:
            query = """
                SELECT geocache_latlng, geocache_address, 
                       geocache_damagetype, geocache_count
                FROM geocache
                ORDER BY geocache_count DESC
                LIMIT 100
            """
            cursor.execute(query)
            rows = cursor.fetchall()
            return rows
    finally:
        conn.close()

def get_subsidence_data():
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password="yhkcar@!0293",
        db="jolupdb",
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor
    )
    try:
        with conn.cursor() as cursor:
            query = """
                SELECT sagoDate, sagoNo, sido, sigungu, dong, addr, sagoDetail
                FROM subsidence_report
                ORDER BY sagoDate DESC
                LIMIT 100
            """
            cursor.execute(query)
            return cursor.fetchall()
    finally:
        conn.close()


# ✅ 루트 접근 안내
@app.get("/")
def root():
    return {
        "message": "🚧 도로 위험 챗봇 서버입니다.",
        "usage": "POST /chat 엔드포인트로 질문을 보내주세요.",
        "example": {
            "url": "/chat",
            "method": "POST",
            "body": {
                "user_question": "고양시에서 은평구로 갈 때 위험한 도로 있어?"
            }
        }
    }

# ✅ GET /chat 안내용
@app.get("/chat")
def chat_info():
    return {
        "message": "❗이 엔드포인트는 POST 요청만 지원합니다.",
        "example_post": {
            "url": "/chat",
            "method": "POST",
            "body": {
                "user_question": "고양시에서 은평구로 갈 때 위험한 도로 있어?"
            }
        }
    }

# ✅ POST /chat: 질문 → Gemini → 응답 반환
@app.post("/chat")
async def chat_with_gemini(question: Question):
    # 1. 데이터 조회
    geocache_data = get_geocache_data()
    subsidence_data = get_subsidence_data()

    # 2. 프롬프트 구성
    prompt = f"""
        당신은 도로 위험 상황을 분석하는 AI 챗봇입니다.
        아래는 두 종류의 실제 데이터입니다.

        📍 [1] 도로 파손 신고 목록 (최근 100건):
        {json.dumps(geocache_data, indent=2)}

        📍 [2] 지반 침하 사고 정보 (최근 100건):
        {json.dumps(subsidence_data, indent=2)}

        🧑 사용자 질문:
        {question.user_question}

        👉 답변 시 아래 지침을 따르세요:
        - 좌표 대신 도로명 주소나 지역명을 사용하세요.
        - 두 데이터 모두 참고하되, 질문 주제에 따라 중요한 데이터를 강조하세요.
        - 정중하고 신뢰감 있는 문장으로 응답하세요.
    """

    # 3. Gemini 응답 생성
    response = model.generate_content(prompt)

    # 4. 응답 반환
    return {"response": response.text}


def summarize_geocache(rows):
    # rows: [{geocache_latlng, geocache_address, geocache_damagetype, geocache_count}, ...]
    total = len(rows)
    by_type = Counter(r["geocache_damagetype"] for r in rows if r.get("geocache_damagetype"))
    # 신고 빈도 상위 주소
    top_addresses = sorted(rows, key=lambda r: r.get("geocache_count", 0), reverse=True)[:10]
    return {
        "total": total,
        "by_type": by_type,
        "top_addresses": top_addresses,
    }

def summarize_subsidence(rows):
    # rows: [{sagoDate, sido, sigungu, dong, addr, ...}, ...]
    total = len(rows)
    by_region = Counter(
        f'{(r.get("sido") or "").strip()} {(r.get("sigungu") or "").strip()}'.strip()
        for r in rows
    )
    # 최근 30일 건수
    recent_cut = (datetime.now()).strftime("%Y%m%d")
    last30 = 0
    try:
        today = datetime.now()
        for r in rows:
            sd = r.get("sagoDate")
            if not sd: 
                continue
            # sagoDate 형식 가정: YYYYMMDD
            d = datetime.strptime(str(sd), "%Y%m%d")
            if (today - d).days <= 30:
                last30 += 1
    except Exception:
        pass
    return {
        "total": total,
        "by_region": by_region,
        "last30": last30,
    }

def build_docx_report(geocache_summary, subsidence_summary, user_note: str = "") -> bytes:
    doc = Document()

    # 제목
    doc.add_heading("도로씨 통계 보고서", 0)
    p = doc.add_paragraph()
    p.add_run(f"생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    if user_note:
        doc.add_paragraph()
        doc.add_paragraph(f"요청 메모: {user_note}")

    # 섹션 1: 도로 파손 신고(Geocache)
    doc.add_heading("1. 도로 파손 신고 요약", level=1)
    doc.add_paragraph(f"총 건수: {geocache_summary['total']}")

    # 유형별 집계 표
    doc.add_paragraph("유형별 건수")
    tbl = doc.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    hdr[0].text = "유형"
    hdr[1].text = "건수"
    for t, c in geocache_summary["by_type"].most_common():
        row = tbl.add_row().cells
        row[0].text = t or "-"
        row[1].text = str(c)

    # 상위 주소
    doc.add_paragraph()
    doc.add_paragraph("신고 빈도 상위 10개 주소")
    tbl2 = doc.add_table(rows=1, cols=3)
    h2 = tbl2.rows[0].cells
    h2[0].text = "주소"
    h2[1].text = "유형"
    h2[2].text = "신고 빈도"
    for r in geocache_summary["top_addresses"]:
        row = tbl2.add_row().cells
        row[0].text = (r.get("geocache_address") or "-")
        row[1].text = (r.get("geocache_damagetype") or "-")
        row[2].text = str(r.get("geocache_count") or 0)

    # 섹션 2: 지반침하(Subsidence)
    doc.add_paragraph()
    doc.add_heading("2. 지반 침하 사고 요약", level=1)
    doc.add_paragraph(f"총 건수: {subsidence_summary['total']}")
    doc.add_paragraph(f"최근 30일 발생 건수: {subsidence_summary['last30']}")

    doc.add_paragraph("지역별 건수(상위 10)")
    tbl3 = doc.add_table(rows=1, cols=2)
    h3 = tbl3.rows[0].cells
    h3[0].text = "지역(시/군/구)"
    h3[1].text = "건수"
    for region, c in subsidence_summary["by_region"].most_common(10):
        row = tbl3.add_row().cells
        row[0].text = region or "-"
        row[1].text = str(c)

    # 바닥글
    doc.add_paragraph()
    doc.add_paragraph("※ 본 보고서는 도로씨 내부 DB(geocache, subsidence_report) 최근 데이터 기준 자동 생성되었습니다.")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# GET: /report.docx?note=발표본
@app.get("/report.docx")
def download_report(note: str = ""):
    geocache_data = get_geocache_data()
    subsidence_data = get_subsidence_data()

    gsum = summarize_geocache(geocache_data)
    ssum = summarize_subsidence(subsidence_data)
    payload = build_docx_report(gsum, ssum, user_note=note)

    filename = f"doro-see-report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        BytesIO(payload),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# POST: /report.docx  (body: {"note": "문구"})
class ReportNote(BaseModel):
    note: str = ""

@app.post("/report.docx")
def download_report_post(req: ReportNote):
    geocache_data = get_geocache_data()
    subsidence_data = get_subsidence_data()

    gsum = summarize_geocache(geocache_data)
    ssum = summarize_subsidence(subsidence_data)
    payload = build_docx_report(gsum, ssum, user_note=req.note)

    filename = f"doro-see-report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        BytesIO(payload),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
